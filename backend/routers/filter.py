from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from backend.database import get_db
from backend.models import Draft, ComplianceReport, ResearchSession
from backend.services.compliance_filter import ComplianceFilter
from backend.services.compliance_report import ComplianceReportGenerator
from docx import Document as DocxDocument
import os
import json

router = APIRouter(prefix="/filter", tags=["Compliance"])


@router.post("/run/{session_id}")
def run_filter(session_id: int, db: Session = Depends(get_db)):

    draft = db.query(Draft).filter(Draft.session_id == session_id).order_by(Draft.id.desc()).first()

    if not draft:
        raise HTTPException(status_code=404, detail="No draft found")

    bias_flags = ComplianceFilter.bias_check(draft.content)
    missing_citations = ComplianceFilter.citation_check(draft.content, db, session_id)
    harmful_flags = ComplianceFilter.harmful_content_check(draft.content)

    report_data = ComplianceReportGenerator.generate(
        bias_flags,
        missing_citations,
        harmful_flags
    )

    report = ComplianceReport(
        session_id=session_id,
        report_json=json.dumps(report_data),
        overall_score=report_data["overall_score"],
        status="PASS" if report_data["passed"] else "FAIL"
    )

    db.add(report)

    if report_data["passed"]:
        session = db.query(ResearchSession).filter(ResearchSession.id == session_id).first()
        session.status = "COMPLETE"

    db.commit()

    return report_data


@router.get("/report/{session_id}")
def get_report(session_id: int, db: Session = Depends(get_db)):
    report = db.query(ComplianceReport).filter(
        ComplianceReport.session_id == session_id
    ).order_by(ComplianceReport.id.desc()).first()

    if not report:
        raise HTTPException(status_code=404, detail="No report found")

    return json.loads(report.report_json)


@router.get("/download/{session_id}")
def download_doc(session_id: int, db: Session = Depends(get_db)):

    draft = db.query(Draft).filter(Draft.session_id == session_id).order_by(Draft.id.desc()).first()

    if not draft:
        raise HTTPException(status_code=404, detail="No draft found")

    file_path = f"uploads/session_{session_id}.docx"
    os.makedirs("uploads", exist_ok=True)

    doc = DocxDocument()
    doc.add_heading("GuardGen Final Draft", level=1)
    doc.add_paragraph(draft.content)
    doc.save(file_path)

    return {"file_saved_at": file_path}